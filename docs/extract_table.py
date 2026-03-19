import sys
try:
    import docx
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

import sys

def docx_table_to_markdown(docx_path):
    try:
        doc = docx.Document(docx_path)
    except Exception as e:
        print(f"Failed to load document at {docx_path}: {e}")
        return

    if not doc.tables:
        print(f"No tables found in {docx_path}")
        return

    # Assume we want the first table, or we process all of them
    for i, table in enumerate(doc.tables):
        print(f"### Table {i+1}\n")
        
        for row_idx, row in enumerate(table.rows):
            row_data = [cell.text.replace('\n', '<br>').strip() for cell in row.cells]
            # Print row
            print("| " + " | ".join(row_data) + " |")
            
            # Print separator after header
            if row_idx == 0:
                separator = ["---"] * len(row.cells)
                print("| " + " | ".join(separator) + " |")
        
        print("\n")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        docx_table_to_markdown(sys.argv[1])
    else:
        print("Please provide the path to the DOCX file.")
